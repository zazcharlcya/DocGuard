import os

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file
)

import docx as python_docx
import pdfplumber

from models import db, Document, ScanResult

from classifier.naive_bayes_classifier import NaiveBayesClassifier
from scanner.pipeline.scanner_pipeline import ScannerPipeline
from scanner.services.file_scanner import FileScanner
from scanner.scanners.email_scanner import EmailScanner
from scanner.scanners.card_scanner import CardScanner
from scanner.scanners.keyword_scanner import KeywordScanner
from scanner.scanners.phone_scanner import PhoneScanner


app = Flask(__name__)

ALLOWED_EXTENSIONS = {"txt", "docx", "pdf"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file_path: str) -> str:
    """Извлекает текст из TXT, DOCX или PDF."""
    ext = file_path.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        doc = python_docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == "pdf":
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n".join(pages)
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


UPLOAD_FOLDER = "uploads"
PROTECTED_UPLOAD_FOLDER = os.path.join(UPLOAD_FOLDER, "protected")
REGULAR_UPLOAD_FOLDER   = os.path.join(UPLOAD_FOLDER, "regular")

for folder in [UPLOAD_FOLDER, PROTECTED_UPLOAD_FOLDER, REGULAR_UPLOAD_FOLDER]:
    os.makedirs(folder, exist_ok=True)

basedir = os.path.abspath(os.path.dirname(__file__))
app.config["SQLALCHEMY_DATABASE_URI"] = (
    "sqlite:///" + os.path.join(basedir, "database", "app.db")
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db.init_app(app)

pipeline = ScannerPipeline()
pipeline.add_scanner(EmailScanner())
pipeline.add_scanner(CardScanner())
pipeline.add_scanner(KeywordScanner())
pipeline.add_scanner(PhoneScanner())

scanner = FileScanner(pipeline)


# ── Дашборд ──────────────────────────────────────────────────
@app.route("/")
def index():
    total        = ScanResult.query.count()
    confidential = ScanResult.query.filter_by(predicted_type="защита").count()
    regular      = ScanResult.query.filter_by(predicted_type="обычный").count()

    stats = {
        "total": total,
        "confidential": confidential,
        "regular": regular,
    }
    return render_template("index.html", stats=stats)


# ── Сканирование ──────────────────────────────────────────────
@app.route("/scan", methods=["POST"])
def scan_file():
    uploaded_file = request.files.get("file")
    if not uploaded_file or uploaded_file.filename == "":
        return "Файл не выбран", 400

    if not allowed_file(uploaded_file.filename):
        return "Неподдерживаемый формат. Допустимы: .txt, .docx, .pdf", 400

    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(file_path)

    file_text = extract_text(file_path)

    document = Document(
        text=file_text,
        type="scan",
        filename=uploaded_file.filename,
        file_path=file_path
    )
    db.session.add(document)
    db.session.commit()

    results = scanner.scan_file(file_path)

    emails, cards, keywords, phones = [], [], [], []

    for result in results:
        if hasattr(result, "emails"):
            emails.extend(result.emails)
        if hasattr(result, "cards"):
            cards.extend(result.cards)
        if hasattr(result, "keywords"):
            keywords.extend(result.keywords)
        if hasattr(result, "phones"):
            phones.extend(result.phones)

    classification    = None
    protection_reasons = []

    if emails:
        protection_reasons.append("найдены email-адреса")
    if cards:
        protection_reasons.append("найдены номера банковских карт")
    if keywords:
        protection_reasons.append("найдены ключевые слова")
    if phones:
        protection_reasons.append("найдены телефонные номера")

    if protection_reasons:
        classification = {
            "predicted_type": "защита",
            "method": "rules",
            "probabilities": {"защита": 1.0, "обычный": 0.0},
            "reasons": protection_reasons
        }
    else:
        training_documents = Document.query.filter(
            Document.type.in_(["защита", "обычный"])
        ).all()

        available_types = {d.type for d in training_documents}

        if "защита" in available_types and "обычный" in available_types:
            classifier = NaiveBayesClassifier()
            classifier.train(training_documents)
            classification = classifier.predict(file_text)
            classification["method"]  = "bayes"
            classification["reasons"] = [
                "чувствительные данные не найдены",
                "использован байесовский классификатор"
            ]

    predicted_type = classification.get("predicted_type") if classification else None

    scan_result = ScanResult(
        document_id=document.id,
        emails=",".join(emails),
        cards=",".join(cards),
        phones=",".join(phones),
        keywords=",".join(keywords),
        predicted_type=predicted_type
    )
    db.session.add(scan_result)
    db.session.commit()

    return render_template(
        "result.html",
        filename=uploaded_file.filename,
        emails=emails,
        cards=cards,
        phones=phones,
        keywords=keywords,
        classification=classification
    )


# ── Защищаемые документы ─────────────────────────────────────
@app.route("/protected")
def protected_documents():
    documents = Document.query.filter_by(type="защита").all()
    return render_template("protected_documents.html", documents=documents)


@app.route("/protected/upload", methods=["GET", "POST"])
def upload_protected_document():
    if request.method == "GET":
        return render_template("upload_protected_document.html")

    uploaded_file = request.files["file"]
    if not allowed_file(uploaded_file.filename):
        return "Неподдерживаемый формат. Допустимы: .txt, .docx, .pdf", 400

    path = os.path.join(PROTECTED_UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(path)

    text = extract_text(path)

    db.session.add(Document(text=text, type="защита",
                            filename=uploaded_file.filename, file_path=path))
    db.session.commit()
    return redirect(url_for("protected_documents"))


@app.route("/protected/<id>/delete", methods=["POST"])
def delete_protected_document(id):
    document = Document.query.get_or_404(id)
    db.session.delete(document)
    db.session.commit()
    return redirect(url_for("protected_documents"))


# ── Обычные документы ────────────────────────────────────────
@app.route("/regular")
def regular_documents():
    documents = Document.query.filter_by(type="обычный").all()
    return render_template("regular_documents.html", documents=documents)


@app.route("/regular/upload", methods=["GET", "POST"])
def upload_regular_document():
    if request.method == "GET":
        return render_template("upload_regular_document.html")

    uploaded_file = request.files["file"]
    if not allowed_file(uploaded_file.filename):
        return "Неподдерживаемый формат. Допустимы: .txt, .docx, .pdf", 400

    path = os.path.join(REGULAR_UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(path)

    text = extract_text(path)

    db.session.add(Document(text=text, type="обычный",
                            filename=uploaded_file.filename, file_path=path))
    db.session.commit()
    return redirect(url_for("regular_documents"))


@app.route("/regular/<id>/delete", methods=["POST"])
def delete_regular_document(id):
    document = Document.query.get_or_404(id)
    db.session.delete(document)
    db.session.commit()
    return redirect(url_for("regular_documents"))


# ── История ──────────────────────────────────────────────────
@app.route("/history")
def scan_history():
    scan_results = ScanResult.query.order_by(ScanResult.created_at.desc()).all()
    return render_template("history.html", scan_results=scan_results)


@app.route("/history/<int:result_id>/delete", methods=["POST"])
def delete_scan_history(result_id):
    scan_result = ScanResult.query.get_or_404(result_id)
    document = scan_result.document
    db.session.delete(scan_result)
    if document and document.type == "scan":
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
        db.session.delete(document)
    db.session.commit()
    return redirect(url_for("scan_history"))


@app.route("/history/<int:result_id>/view")
def view_scan_result(result_id):
    scan_result = ScanResult.query.get_or_404(result_id)
    document    = scan_result.document
    emails   = [e for e in (scan_result.emails   or "").split(",") if e]
    cards    = [c for c in (scan_result.cards    or "").split(",") if c]
    phones   = [p for p in (scan_result.phones   or "").split(",") if p]
    keywords = [k for k in (scan_result.keywords or "").split(",") if k]
    return render_template(
        "history_detail.html",
        scan_result=scan_result,
        document=document,
        emails=emails,
        cards=cards,
        phones=phones,
        keywords=keywords,
    )


@app.route("/history/<int:result_id>/download")
def download_scan_report(result_id):
    scan_result = ScanResult.query.get_or_404(result_id)
    document = scan_result.document

    report_text = f"""Отчёт проверки документа
========================================
Файл:              {document.filename if document else "неизвестно"}
Дата:              {scan_result.created_at}
Категория:         {scan_result.predicted_type or "не определена"}

Email-адреса:
{scan_result.emails or "не найдено"}

Телефонные номера:
{scan_result.phones or "не найдено"}

Номера банковских карт:
{scan_result.cards or "не найдено"}

Ключевые слова:
{scan_result.keywords or "не найдено"}
"""

    reports_folder = os.path.join(UPLOAD_FOLDER, "reports")
    os.makedirs(reports_folder, exist_ok=True)
    report_path = os.path.join(reports_folder, f"report_{scan_result.id}.txt")

    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_text)

    return send_file(report_path, as_attachment=True,
                     download_name=f"report_{scan_result.id}.txt")



if __name__ == "__main__":
    with app.app_context():
        # Создаём папку database и таблицы при первом запуске
        os.makedirs(os.path.join(basedir, "database"), exist_ok=True)
        db.create_all()

        # Автомиграция: добавляем колонку phones если её нет (обновление со старой версии)
        from sqlalchemy import text, inspect as sa_inspect
        try:
            insp = sa_inspect(db.engine)
            existing_cols = [c["name"] for c in insp.get_columns("scan_result")]
            if "phones" not in existing_cols:
                with db.engine.connect() as conn:
                    conn.execute(text("ALTER TABLE scan_result ADD COLUMN phones TEXT"))
                    conn.commit()
                print("[INFO] Добавлена колонка phones в scan_result")
        except Exception:
            pass

    app.run(debug=True)
